#!/usr/bin/env python3
import random
import re
import argparse
from pathlib import Path
from datetime import datetime

# Optional deps
try:
    from docx import Document
except Exception:
    Document = None
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
except Exception:
    canvas = None


SAMPLE_DIR = Path(__file__).resolve().parents[1] / "Sample_files"


SOP_TITLES = [
    "PBMC Isolation Protocol",
    "Flow Cytometry Panel Staining",
    "Cell Culture Maintenance",
    "RNA Extraction and QC",
    "ELISA Quantitation Assay",
    "Western Blot Detection",
    "Immunofluorescence Imaging",
    "Mouse Dosing and Sample Collection",
    "qPCR Expression Analysis",
    "Cryopreservation of PBMCs",
    "T Cell Activation Assay",
]


def make_docx(path: Path, title: str, body: str):
    if not Document:
        raise RuntimeError("python-docx is required to generate .docx files")
    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(str(path))


def make_pdf(path: Path, title: str, paragraphs: list[str]):
    if not canvas:
        raise RuntimeError("reportlab is required to generate PDFs")
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 1 * inch
    c.setFont("Helvetica-Bold", 16)
    c.drawString(1 * inch, y, title)
    y -= 0.4 * inch
    c.setFont("Helvetica", 10)
    for para in paragraphs:
        for line in wrap_text(para, 90):
            if y < 1 * inch:
                c.showPage()
                y = height - 1 * inch
                c.setFont("Helvetica", 10)
            c.drawString(1 * inch, y, line)
            y -= 0.18 * inch
        y -= 0.12 * inch
    c.showPage()
    c.save()


def wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    lines = []
    cur = []
    count = 0
    for w in words:
        if count + len(w) + (1 if cur else 0) > width:
            lines.append(" ".join(cur))
            cur = [w]
            count = len(w)
        else:
            cur.append(w)
            count += len(w) + (1 if cur else 0)
    if cur:
        lines.append(" ".join(cur))
    return lines


def generate_sop_text(title: str) -> str:
    steps = [
        "Purpose: This SOP describes standardized procedures to ensure data integrity and reproducibility.",
        "Materials: List all reagents with catalog numbers, instruments (e.g., BD LSRFortessa), and consumables.",
        "Safety: Follow BSL-2 practices and institutional IACUC/IRB approvals.",
        "Procedure: Detail volume, incubation time, temperature, centrifugation speed, and gating strategy.",
        "QC: Include acceptance criteria, positive/negative controls, and repeat thresholds.",
        "Documentation: Record lot numbers, deviations, and corrective actions in ELN.",
    ]
    random.shuffle(steps)
    return f"{title}\n\n" + "\n\n".join(steps)


def generate_rfp_paragraphs(n: int = 3) -> list[str]:
    parts = [
        "Background: Sponsor requests a preclinical immunogenicity study assessing T cell responses to candidate antigen.",
        "Scope: Include PBMC isolation, flow cytometry panel with 16 colors (CD3, CD4, CD8, CD45RA, CCR7, IFNγ, TNFα, IL-2, etc.), and ELISA for cytokines.",
        "Data: Provide sample size justification (power=0.8), expected effect sizes, and statistical plan (ANOVA with post-hoc tests).",
        "Deliverables: Raw FCS files, gating strategy report, annotated spreadsheets, and a final PDF summary.",
        "Timeline: 6 weeks from sample receipt; milestones at week 2 and week 4.",
        "Budget: Provide line-item costs for reagents, instrument time, analyst hours, and QA review.",
    ]
    random.shuffle(parts)
    return parts[:n]


# --- Enhancements: Customers and Pricing ---
CUSTOMERS = ["Pfizer", "BMS", "Moderna", "J&J"]

LINE_ITEM_NAMES = [
    "Flow cytometry panel (16-color)",
    "ELISA cytokine panel",
    "Animal housing (per cage)",
    "Scientist analysis (per hour)",
    "QA review",
    "PBMC isolation (per sample)",
    "qPCR reagents",
    "Western blot reagents",
    "Microscopy imaging time",
]

def generate_pricing_items(min_items: int = 3, max_items: int = 6) -> tuple[list[dict], float]:
    k = random.randint(min_items, max_items)
    picks = random.sample(LINE_ITEM_NAMES, k)
    items: list[dict] = []
    subtotal = 0.0
    for name in picks:
        qty = random.randint(1, 12)
        unit_cost = round(random.uniform(80, 2500), 2)
        line_total = round(qty * unit_cost, 2)
        subtotal += line_total
        items.append({
            "item": name,
            "quantity": qty,
            "unit_cost": unit_cost,
            "line_total": line_total,
        })
    subtotal = round(subtotal, 2)
    return items, subtotal

def format_money(amount: float) -> str:
    return f"${amount:,.2f}"

def build_customer_pricing_paragraphs() -> list[str]:
    customer = random.choice(CUSTOMERS)
    items, subtotal = generate_pricing_items()
    lines = [f"Customer: {customer}", "Pricing:"]
    for it in items:
        lines.append(
            f"- {it['item']} — qty {it['quantity']} × {format_money(it['unit_cost'])} = {format_money(it['line_total'])}"
        )
    lines.append(f"Subtotal: {format_money(subtotal)}")
    return ["\n".join(lines)]

def next_sop_index() -> int:
    max_n = 0
    pat = re.compile(r"SOP_(\d{2,})\.(?:pdf|docx)$", re.I)
    for p in SAMPLE_DIR.glob("SOP_*.*"):
        m = pat.search(p.name)
        if m:
            try:
                n = int(m.group(1))
                max_n = max(max_n, n)
            except Exception:
                pass
    return max_n + 1

def generate_extra_sops(count: int = 15) -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    start = next_sop_index()
    for i in range(count):
        idx = start + i
        title = random.choice(SOP_TITLES)
        # Build paragraphs: SOP body + customer + pricing
        body_paras = generate_sop_text(title).split("\n\n")
        body_paras.extend(build_customer_pricing_paragraphs())
        # Mix of DOCX/PDF; alternate for even distribution
        use_docx = (i % 2 == 0)
        if use_docx:
            out = SAMPLE_DIR / f"SOP_{idx:02d}.docx"
            make_docx(out, title, "\n\n".join(body_paras))
        else:
            out = SAMPLE_DIR / f"SOP_{idx:02d}.pdf"
            make_pdf(out, title, body_paras)


def main():
    parser = argparse.ArgumentParser(description="Generate sample SOPs/RFPs")
    parser.add_argument("--extra-sops", type=int, default=0, help="Generate this many additional SOPs (mixed DOCX/PDF) with Customer and Pricing sections")
    parser.add_argument("--all", action="store_true", help="Generate the original baseline set (10 SOP docx + 10 SOP pdf + 10 RFP pdf)")
    args = parser.parse_args()

    if args.extra_sops and args.extra_sops > 0:
        generate_extra_sops(args.extra_sops)
        print(f"Generated {args.extra_sops} additional SOPs in {SAMPLE_DIR}")
        return

    # Default: original baseline set
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)

    # Generate 10 SOP DOCX
    for i in range(10):
        title = random.choice(SOP_TITLES)
        body = generate_sop_text(title)
        out = SAMPLE_DIR / f"SOP_{i+1:02d}.docx"
        make_docx(out, title, body)

    # Generate 10 SOP PDFs
    for i in range(10):
        title = random.choice(SOP_TITLES)
        paras = generate_sop_text(title).split("\n\n")
        out = SAMPLE_DIR / f"SOP_{i+11:02d}.pdf"
        make_pdf(out, title, paras)

    # Generate 10 RFP PDFs with scientific data text
    for i in range(10):
        title = f"Request for Proposal (RFP) #{i+1:02d} — {datetime.utcnow().date()}"
        paras = generate_rfp_paragraphs(4)
        out = SAMPLE_DIR / f"RFP_{i+1:02d}.pdf"
        make_pdf(out, title, paras)

    print(f"Generated samples in {SAMPLE_DIR}")


if __name__ == "__main__":
    main()


