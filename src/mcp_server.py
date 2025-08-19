from fastapi import FastAPI, UploadFile, File, Form, Request, Depends, Header
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi_mcp import FastApiMCP
from typing import Dict, Optional
from contextvars import ContextVar
from time import time
from datetime import datetime
from pydantic import BaseModel, Field
from src.milvus_connector import search_similar_documents, add_document, ensure_collection_exists
from src.sow_generator import generate_sow
import asyncio
from pathlib import Path
from src.constants import mcp_config
from src.constants import (
    CLERK_PUBLISHABLE_KEY,
    CLERK_SECRET_KEY,
    CLERK_FRONTEND_API,
    CLERK_API_URL,
    CLERK_JWKS_URL,
    INTERNAL_API_SECRET,
)
from mcp_use import MCPClient, MCPAgent
from langchain_openai import ChatOpenAI
from src.order_store import create_order, list_orders, update_status, get_order, update_items, update_customer
from src.audit_log import append_log, list_logs
from src.mcp_client import main as mcp_client_run
import re
import json as _json
# Clerk disabled for now
Clerk = None

try:
    from docx import Document
except Exception:
    Document = None
from io import BytesIO
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

# Optional PDF generation dependency
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
except Exception:
    canvas = None

app = FastAPI()

# Lightweight singletons to reduce per-request initialization cost
_MCP_CLIENT: MCPClient | None = None
_LLM: ChatOpenAI | None = None

def _get_mcp_client() -> MCPClient:
    global _MCP_CLIENT
    if _MCP_CLIENT is None:
        _MCP_CLIENT = MCPClient(mcp_config)
    return _MCP_CLIENT

def _get_llm() -> ChatOpenAI:
    global _LLM
    if _LLM is None:
        _LLM = ChatOpenAI(model="gpt-4o-mini", streaming=False)
    return _LLM

# Optional Clerk SDK initialization (no-op if package/env not configured yet)
def _init_clerk() -> Optional[Clerk]:
    try:
        if Clerk is None:
            return None
        # Use constants loaded from env
        secret = CLERK_SECRET_KEY
        if not secret:
            return None
        return Clerk(secret_key=secret)
    except Exception:
        return None

# _clerk = _init_clerk()

# Auth utilities (JWT via Clerk JWKS)
# JWT auth disabled for now
jwt = None
PyJWKClient = None

def _auth_enabled() -> bool:
    return False

_jwks_client = None
def _get_jwks_client():
    global _jwks_client
    if _jwks_client is None and _auth_enabled():
        try:
            _jwks_client = PyJWKClient(CLERK_JWKS_URL)
        except Exception:
            _jwks_client = None
    return _jwks_client

async def require_auth(authorization: str | None = Header(default=None), clerk_session: str | None = Header(default=None), x_internal_secret: str | None = Header(default=None)) -> Optional[dict]:
    # Auth temporarily disabled
    return None
# Middleware to log orders-related API calls with prompt and tool_name (operation_id)
@app.middleware("http")
async def audit_orders_calls(request: Request, call_next):
    response = await call_next(request)
    try:
        path = request.url.path or ""
        method = request.method
        # Only audit orders endpoints here (write prompt/tool_name). Others already log individually.
        if path.startswith("/orders"):
            # Try to resolve operation_id as tool_name
            tool = None
            try:
                route = request.scope.get("route")
                tool = getattr(route, "operation_id", None) or getattr(route, "name", None)
            except Exception:
                tool = None
            append_log({
                "type": "api_call",
                "route": path,
                "method": method,
                "status": getattr(response, "status_code", None),
                # Summarize recent user prompt; never log tool guide
                "prompt": _audit_prompt_context(),
                "tool_name": tool,
            })
    except Exception:
        pass
    return response


# Context variable to carry the active agent prompt across tool calls
AGENT_PROMPT: ContextVar[str | None] = ContextVar("AGENT_PROMPT", default=None)
AGENT_ORDER_ID: ContextVar[str | None] = ContextVar("AGENT_ORDER_ID", default=None)
def _fmt_money(amount: float | int) -> str:
    try:
        return f"${float(amount):,.0f}"
    except Exception:
        return "$0"

def _shorten(text: str | None, limit: int = 200) -> str:
    s = (text or "").strip()
    return s[:limit] + ("…" if len(s) > limit else "")

def _summarize_text(text: str | None, limit: int = 200) -> str:
    if not text:
        return ""
    s = str(text).strip()
    # Prefer the first line/sentence for quick context
    first_line = s.splitlines()[0] if "\n" in s else s
    # Try to break at sentence end
    for sep in [". ", "! ", "? "]:
        if sep in first_line and len(first_line.split(sep)[0]) >= 32:
            first_line = first_line.split(sep)[0] + sep.strip()
            break
    return _shorten(first_line, limit)

def _summarize_additions(additions: list) -> str:
    try:
        additions = additions or []
        if not additions:
            return "no extra items"
        if len(additions) == 1:
            it = additions[0]
            name = str(it.get("item", "item")).strip() or "item"
            cost = float(it.get("unit_cost", it.get("cost", 0)) or 0)
            return f"added 1 item: {name} {_fmt_money(cost)}"
        names = ", ".join(str((i.get("item") or "item")).strip() for i in additions[:3])
        more = "" if len(additions) <= 3 else f" (+{len(additions)-3} more)"
        return f"added {len(additions)} items: {names}{more}"
    except Exception:
        return "added items"

def _summarize_items(items: list) -> str:
    try:
        items = items or []
        subtotal = 0.0
        for it in items:
            q = float(it.get("quantity", 0) or 0)
            c = float(it.get("unit_cost", it.get("cost", 0)) or 0)
            subtotal += q * c
        return f"updated {len(items)} items; subtotal {_fmt_money(subtotal)}"
    except Exception:
        return "items updated"

def _summarize_agent_result(result: object) -> str:
    # Try to produce a compact human summary from agent output
    try:
        text = result if isinstance(result, str) else _json.dumps(result)
    except Exception:
        text = str(result)
    # Try JSON
    try:
        data = _json.loads(text) if isinstance(text, str) else result
        if isinstance(data, dict):
            created = data.get("created_id") or (data.get("order") or {}).get("id")
            updated = data.get("updated_ids") or []
            actions = data.get("actions") or []
            notes = data.get("notes")
            parts: list[str] = []
            if created:
                parts.append(f"created order {created}")
            if updated:
                parts.append(f"updated {len(updated)} order(s)")
            if actions:
                parts.append(", ".join(actions))
            if notes and not parts:
                parts.append(str(notes))
            if parts:
                return "; ".join(parts)
    except Exception:
        pass
    # Fallback: extract OD- ids and a short snippet
    try:
        m = re.findall(r"OD-\d{5}", text)
        lead = text.strip().splitlines()[0][:140] if isinstance(text, str) else str(text)[:140]
        if m:
            return f"agent action on {', '.join(sorted(set(m)))}: {lead}"
        return lead
    except Exception:
        return "agent action executed"

# Heuristic extraction for customer name from freeform notes
def _extract_customer_from_notes(notes: str) -> Optional[str]:
    try:
        s = (notes or "").strip()
        if not s:
            return None
        # 1) Structured labels
        patterns = [
            r"(?im)^[\-*\s>]*customer(?:\s+name)?\s*[:\-]\s*(.+)$",
            r"(?im)^[\-*\s>]*client\s*[:\-]\s*(.+)$",
            r"(?im)^[\-*\s>]*company\s*[:\-]\s*(.+)$",
        ]
        for pat in patterns:
            m = re.search(pat, s)
            if m:
                cand = m.group(1).strip()
                cand = re.sub(r"^\**|\**$", "", cand).strip()
                cand = re.sub(r"^`+|`+$", "", cand).strip()
                # Stop at common delimiters and trailing sections like Pricing
                cand = re.split(r"(?:\s[–—-]\s|[\|•#/:;,\r\n]|\s+Pricing:?|\s+PRICING:?|\s+pricing:?)", cand)[0].strip()
                if 1 <= len(cand) <= 120:
                    return cand
        # 2) Conversational hints (e.g., "assume it's Pfizer")
        m2 = re.search(r"(?i)assume(?:\s+it\s+is|\s+it's|\s+its)?\s+([A-Za-z][A-Za-z0-9 .&'-]{1,120})", s)
        if m2:
            cand = m2.group(1).strip()
            cand = re.sub(r"[^A-Za-z0-9 .&'-]", "", cand).strip()
            if 1 <= len(cand) <= 120:
                return cand
        # 3) Known company names (quick dictionary match)
        known = [
            "Pfizer", "Moderna", "BMS", "Bristol Myers", "GSK", "Novartis", "Merck",
            "Roche", "Sanofi", "Amgen", "AstraZeneca", "Regeneron", "Biogen", "Bayer",
            "AbbVie", "Eli Lilly",
        ]
        for name in known:
            if re.search(rf"(?i)\b{name}\b", s):
                return name
        return None
    except Exception:
        return None


def _canonicalize_customer(name: Optional[str]) -> Optional[str]:
    """Map a candidate customer name to an existing known customer label if present.

    Preference order: exact (case-insensitive) match of full string among existing orders; else
    match ignoring punctuation and collapsing spaces; else return input name.
    """
    try:
        if not name:
            return name
        cand = str(name).strip()
        if not cand:
            return name
        # Collect known customers from current orders
        existing = []
        try:
            for o in list_orders():
                c = (o.get("customer") or "").strip()
                if c:
                    existing.append(c)
        except Exception:
            existing = []
        if not existing:
            return name
        def norm(x: str) -> str:
            x = x.lower().strip()
            x = re.sub(r"[^a-z0-9\s&'\.-]", "", x)
            x = re.sub(r"\s+", " ", x)
            return x
        cand_n = norm(cand)
        # Exact case-insensitive
        for k in existing:
            if cand.lower() == k.lower():
                return k
        # Normalized equality
        for k in existing:
            if norm(k) == cand_n:
                return k
        # Substring containment heuristic
        for k in existing:
            nk = norm(k)
            if cand_n in nk or nk in cand_n:
                return k
        return name
    except Exception:
        return name

# Global prompt cache to bridge across separate HTTP requests initiated by the agent
LAST_AGENT_PROMPT: str | None = None
LAST_AGENT_PROMPT_TS: float = 0.0
PROMPT_TTL_SECONDS: float = 300.0
LAST_AGENT_ORDER_ID: str | None = None
LAST_AGENT_ORDER_ID_TS: float = 0.0
LAST_USER_PROMPT: str | None = None
LAST_USER_PROMPT_TS: float = 0.0

# Track last assistant message (e.g., a follow-up question) to improve audit trail context
LAST_ASSISTANT_MSG: str | None = None
LAST_ASSISTANT_MSG_TS: float = 0.0

def remember_agent_prompt(prompt: Optional[str]) -> None:
    global LAST_AGENT_PROMPT, LAST_AGENT_PROMPT_TS
    if prompt:
        LAST_AGENT_PROMPT = str(prompt)
        LAST_AGENT_PROMPT_TS = time()

def recent_agent_prompt() -> Optional[str]:
    try:
        if LAST_AGENT_PROMPT and (time() - LAST_AGENT_PROMPT_TS) <= PROMPT_TTL_SECONDS:
            return LAST_AGENT_PROMPT
    except Exception:
        pass
    return None

def remember_user_prompt(prompt: Optional[str]) -> None:
    global LAST_USER_PROMPT, LAST_USER_PROMPT_TS
    if prompt:
        LAST_USER_PROMPT = str(prompt)
        LAST_USER_PROMPT_TS = time()

def recent_user_prompt() -> Optional[str]:
    try:
        if LAST_USER_PROMPT and (time() - LAST_USER_PROMPT_TS) <= PROMPT_TTL_SECONDS:
            return LAST_USER_PROMPT
    except Exception:
        pass
    return None

def remember_assistant_message(message: Optional[str]) -> None:
    global LAST_ASSISTANT_MSG, LAST_ASSISTANT_MSG_TS
    if message:
        LAST_ASSISTANT_MSG = str(message)
        LAST_ASSISTANT_MSG_TS = time()

def recent_assistant_message() -> Optional[str]:
    try:
        if LAST_ASSISTANT_MSG and (time() - LAST_ASSISTANT_MSG_TS) <= PROMPT_TTL_SECONDS:
            return LAST_ASSISTANT_MSG
    except Exception:
        pass
    return None

def _looks_like_confirmation(text: str | None) -> bool:
    if not text:
        return False
    s = str(text).strip().lower()
    return s in {"yes", "y", "yeah", "yep", "sure", "ok", "okay", "please do", "do it", "proceed"} or s in {"no", "n", "nope", "don't", "do not", "cancel"}

def _audit_prompt_context(fallback: Optional[str] = None) -> str:
    """Return a concise prompt summary for audit logs.

    If the latest user message is a short confirmation (yes/no/sure), include the
    previous assistant follow-up question for clarity: "yes — Follow-up: <question>".
    Otherwise, summarize the latest user prompt. If none, fall back to agent prompt.
    """
    try:
        user = recent_user_prompt()
        if user:
            user_s = _summarize_text(user, 200)
            if _looks_like_confirmation(user_s):
                a = recent_assistant_message()
                if a:
                    return f"{user_s} — Follow-up: {_summarize_text(a, 200)}"
            return user_s
        return _summarize_text(AGENT_PROMPT.get(), 200) or _summarize_text(fallback, 200)
    except Exception:
        return _summarize_text(fallback or recent_user_prompt() or AGENT_PROMPT.get(), 200)

def remember_agent_order_id(order_id: Optional[str]) -> None:
    global LAST_AGENT_ORDER_ID, LAST_AGENT_ORDER_ID_TS
    if order_id:
        LAST_AGENT_ORDER_ID = str(order_id)
        LAST_AGENT_ORDER_ID_TS = time()

def recent_agent_order_id() -> Optional[str]:
    try:
        if LAST_AGENT_ORDER_ID and (time() - LAST_AGENT_ORDER_ID_TS) <= PROMPT_TTL_SECONDS:
            return LAST_AGENT_ORDER_ID
    except Exception:
        pass
    return None

# Order ID normalization helpers
def _extract_order_numeric(order_id: str) -> Optional[int]:
    try:
        s = str(order_id or "").strip()
        # Accept patterns like OD-00023, od-23, OD00023
        import re as _re
        m = _re.search(r"(\d+)$", s)
        if not m:
            return None
        return int(m.group(1))
    except Exception:
        return None

def _resolve_order_id(candidate: str | None) -> Optional[str]:
    if not candidate:
        return None
    target_n = _extract_order_numeric(candidate)
    if target_n is None:
        return None
    try:
        for o in list_orders():
            n = _extract_order_numeric(o.get("id"))
            if n == target_n:
                return o.get("id")
    except Exception:
        pass
    return None

class AddDocumentRequest(BaseModel):
    text: str = Field(..., description="Document text to add")
    metadata: dict = Field(default_factory=dict, description="Optional metadata for the document")


@app.post('/v1/search', operation_id="search_for_documents")
async def search(query: str) -> list:
    """Search for documents in Milvus."""
    return await search_similar_documents(query)


@app.post('/v1/add-documents', operation_id="add_documents_to_milvus")
async def add_document_to_milvus(request: AddDocumentRequest) -> Dict:
    """Add a document to Milvus vector database."""
    try:
        print(f"Adding document: {request.text[:100]}...")
        print(f"Metadata: {request.metadata}")
        
        result = await add_document(request.text, request.metadata)
        
        return {
            "success": True,
            "message": "Document added successfully to Milvus",
            "result": result
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to add document to Milvus: {str(e)}",
            "text": request.text,
            "metadata": request.metadata
        }

# Define a simple tool (disabled)
# @app.get("/v1/greet")#, operation_id="greet")
# async def greet(name: str) -> Dict:
#     """Greets the given name."""
#     return {"message": f"Hello, {name}!"}

# @app.get("/v2/greet", operation_id="greet")
# async def greet(name: str) -> Dict:
#     """Greets the given name."""
#     return {"message": f"Howdy, {name}!"}

# Serve Home page (Agent)
@app.get("/", operation_id="home_page")
async def home_page():
    # Serve agent.html as the home page
    page = static_dir / "agent.html"
    if page.exists():
        return FileResponse(str(page), media_type="text/html")
    return JSONResponse(status_code=404, content={"error": "agent.html not found"})

# NOTE: MCP mounting moved to the end of the file to ensure ALL routes are exposed as tools

# Mount static files for a simple frontend UI
static_dir = Path(__file__).resolve().parent / "static"
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")


@app.post("/sow/upload", operation_id="sow_upload")
async def sow_upload(
    file: UploadFile = File(...),
    customer: Optional[str] = Form(default=None),
    title: Optional[str] = Form(default=None),
):
    """Accept a PDF or DOCX, extract text, generate a SOW, and return it."""
    content = await file.read()
    text = ""
    try:
        if file.filename.lower().endswith(".docx"):
            if not Document:
                return {"error": "python-docx not installed"}
            doc = Document(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        elif file.filename.lower().endswith(".pdf"):
            if not PdfReader:
                return {"error": "pypdf not installed"}
            reader = PdfReader(BytesIO(content))
            pages = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t)
            text = "\n".join(pages)
        else:
            return {"error": "Unsupported file type. Upload .pdf or .docx"}
    except Exception as e:
        return {"error": f"Failed to read file: {e}"}

    if not text.strip():
        return {"error": "No extractable text found in the file"}

    sow = await generate_sow(text, customer_name=customer, project_title=title)
    return {"sow": sow}


@app.get("/sow", operation_id="sow_form")
async def sow_form():
    return {
        "html": """
<!doctype html>
<html>
  <head>
    <meta charset=\"utf-8\" />
    <title>SOW Generator</title>
  </head>
  <body>
    <h1>SOW Generator</h1>
    <form id=\"form\" method=\"post\" action=\"/sow/upload\" enctype=\"multipart/form-data\">
      <label>Customer (optional): <input type=\"text\" name=\"customer\" /></label><br/>
      <label>Title (optional): <input type=\"text\" name=\"title\" /></label><br/>
      <input type=\"file\" name=\"file\" accept=\".pdf,.docx\" required />
      <button type=\"submit\">Generate SOW</button>
    </form>
  </body>
  </html>
        """
    }


class AgentRunRequest(BaseModel):
    prompt: str = Field(..., description="Prompt to run through the MCP agent")
    history: list | None = Field(default=None, description="Optional conversation history: [{role, content}]")
    order_id: Optional[str] = Field(default=None, description="Current order context. Default for tools if user says 'this order'.")


@app.post("/agent/run", operation_id="agent_run")
async def agent_run(req: AgentRunRequest):
    try:
        client = _get_mcp_client()
        llm = _get_llm()
        agent = MCPAgent(llm=llm, client=client, max_steps=20)

        TOOL_GUIDE = (
            "System instructions:\n"
            "- Discover available tools dynamically from the MCP server schema and choose the minimal set of actions to fulfill the user's intent.\n"
            "- Plan before acting: retrieve any required data first (e.g., list entities), reason in-memory (sort/filter/compare), then execute precise updates or creations.\n"
            "- Respect data safety: do not modify a source entity when the user asks to create a similar/variant entity unless they explicitly request modifications to the original.\n"
            "- Prefer high-level tools when available (e.g., clone/variant creators) rather than orchestrating multi-step sequences, but fall back gracefully to lower-level tools.\n"
            "- Ensure idempotency: avoid duplicate submissions; each intended action should execute once.\n"
            "- Handle errors pragmatically: if an action fails (e.g., not found), re-synchronize by listing relevant data and retry once if appropriate; otherwise report succinctly.\n"
            "- Act when tools exist; do not claim lack of access. If no suitable tool exists, explain limitations briefly.\n"
            "- Scope guardrails: Only perform actions that map directly to available MCP tools/API endpoints in this app. Never claim or simulate capabilities that are not exposed here.\n"
            "- Mixed-scope requests: If a request contains both supported and unsupported parts, DO NOT execute any action yet. First, ask ONE concise follow-up to confirm proceeding with the supported part(s), then await explicit user approval (e.g., 'yes') before calling any tools. Begin the reply with 'Follow-up:' so the user understands it's a clarification.\n"
            "- Unsupported examples (decline politely): drawing/creating images, posting to Instagram or other external sites, sending emails/Slack, web browsing, payment processing, user/account administration.\n"
            "  When declining, say: 'We’re sorry, but we can’t do that in this app.' Offer a closest in-scope alternative only if helpful.\n"
            "- If the user asks to export to Excel/CSV or similar external operations, acknowledge limitation and ask a constructive follow-up like: 'Follow-up: I can list or summarize all Pfizer orders, compute totals, or change statuses. What would you like me to do with these orders?'\n"
            "- Clarification & confirmation: If the mapping to a tool is ambiguous or could impact multiple records, ask ONE concise follow-up question and await the user's confirmation before proceeding.\n"
            "- Example: If asked 'create a quote like OD-37 and draw pictures of deer and elephant and upload to my Instagram':\n"
            "  → Ask: 'Follow-up: I can create a new quote like OD-37. Proceed?'\n"
            "  → After user's 'yes', perform clone_order/create_variant_order.\n"
            "  → Also state in the same message that drawing images and Instagram uploads are not supported in this app.\n"
            "- Tone: Be neutral and helpful. Avoid accusatory phrasing (e.g., 'you requested twice'). Prefer phrasing like 'Just to confirm…' or 'Would you like me to…'.\n"
            "- Potential duplicates: If the user asks to add a line item that is similar to an existing one (e.g., another Service Fee $1,500), do NOT assume it's a mistake. Ask a short confirmation like: 'Add another Service Fee of $1,500 as a separate line item?' and wait for 'yes' before proceeding.\n"
            "- When user asks to add a fee or line item to an existing order id, call add_items_to_order with that id. Do NOT clone the order.\n"
            "- When user asks to change/update the customer name for an existing order id, call update_order_customer with that id. Do NOT add a line item for customer.\n"
            "- For 'clone then add items' flows only when user explicitly asks for a variant/new order: call clone_order → (optionally) add_items_to_order.\n"
            "- STRICT FORMAT for add_items_to_order: additions must be a JSON array of objects with exact keys { item: string, quantity: number, unit_cost: number }. Example: {\n"
            "    \"order_id\": \"<CURRENT_ORDER_ID>\",\n"
            "    \"additions\": [{ \"item\": \"Scientist Flat Fee\", \"quantity\": 10, \"unit_cost\": 3500 }]\n"
            "  } . Do NOT send free-text or nested strings. You may omit line_total; the server computes it.\n"
            "- Accept synonyms in user input (qty, q, price, unit price), but ALWAYS send the API fields as quantity and unit_cost in the final tool call.\n"
            "- Omit optional fields when not set. Do NOT send null for optional string fields (e.g., omit 'prompt' instead of sending null).\n"
            f"- Current order context: {req.order_id or '(none provided)'} . When the user says 'this order' or omits order_id for order tools, use this order id by default.\n"
            "- For PDF preview of an existing order, call orders_pdf (GET /orders/pdf/{order_id}). Only call agent_quote_pdf when you provide a body with items and subtotal.\n"
            "- Output style: Write for end users in brief, readable sentences or 2–6 short bullets. Avoid code blocks and JSON unless explicitly requested.\n"
            "  Present key figures clearly (e.g., Today total: $X; Yesterday total: $Y), then 1–3 bullets with highlights. Use US currency formatting.\n"
        )
        # Build conversational context from history
        convo = []
        if req.history:
            for msg in req.history:
                role = str(msg.get("role", "")).strip().lower()
                content = str(msg.get("content", ""))
                if role == "assistant":
                    convo.append(f"Assistant: {content}")
                else:
                    convo.append(f"User: {content}")
        convo.append(f"User: {req.prompt}")
        combined_prompt = f"{TOOL_GUIDE}\n\n" + "\n".join(convo)

        # Set prompt context for downstream tool calls
        prompt_token = AGENT_PROMPT.set(req.prompt)
        order_token = None
        if req.order_id:
            order_token = AGENT_ORDER_ID.set(req.order_id)
        # Persist full reasoning context (tool guide + conversation) for downstream audit logging
        remember_agent_prompt(combined_prompt)
        remember_user_prompt(req.prompt)
        remember_agent_order_id(req.order_id)

        # Log start with a concise summary of the user prompt (not the tool guide)
        user_prompt_summary = _summarize_text(req.prompt, 200)
        # Log summarized user prompt only (exclude tool guide)
        append_log({"type": "agent_run_start", "prompt": user_prompt_summary})

        result = await agent.run(combined_prompt)
        # Keep a concise summary of the user prompt for audit trail, do not log tool guide
        action_summary = _summarize_text(req.prompt, 200)
        append_log({
            "type": "agent_run",
            # Use audit prompt context to capture follow-up + yes/no
            "prompt": _audit_prompt_context(action_summary),
            "tokens_output": int(max(1, len(str(result))//4)),
            "result": result,
        })
        return {"result": result}
    except Exception as e:
        # On error, prefer summarizing the user prompt
        safe_prompt = _summarize_text(req.prompt if 'req' in locals() else "", 200)
        append_log({"type": "agent_run_error", "prompt": safe_prompt, "error": str(e)})
        return JSONResponse(status_code=400, content={"error": str(e)})
    finally:
        try:
            # Reset context to avoid leaking prompt across requests
            AGENT_PROMPT.reset(prompt_token)
        except Exception:
            pass
        try:
            if 'order_token' in locals() and order_token is not None:
                AGENT_ORDER_ID.reset(order_token)
        except Exception:
            pass


@app.get("/cpq", operation_id="cpq_page")
async def cpq_page():
    page = static_dir / "cpq.html"
    if page.exists():
        return FileResponse(str(page), media_type="text/html")
    return JSONResponse(status_code=404, content={"error": "cpq.html not found"})


# @app.get("/rfps")
# async def rfps_page():
#     page = static_dir / "rfps.html"
#     if page.exists():
#         return FileResponse(str(page), media_type="text/html")
#     return JSONResponse(status_code=404, content={"error": "rfps.html not found"})


@app.get("/notetaker", operation_id="notetaker_page")
async def notetaker_page():
    page = static_dir / "notetaker.html"
    if page.exists():
        return FileResponse(
            str(page),
            media_type="text/html",
            headers={
                "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
                "Pragma": "no-cache",
            },
        )
    return JSONResponse(status_code=404, content={"error": "notetaker.html not found"})


@app.get("/orders", operation_id="orders_page")
async def orders_page():
    page = static_dir / "orders.html"
    if page.exists():
        return FileResponse(str(page), media_type="text/html")
    return JSONResponse(status_code=404, content={"error": "orders.html not found"})


def _chunk_text(text: str, chunk_size: int = 1200, chunk_overlap: int = 200):
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = max(end - chunk_overlap, start + 1)
    return chunks


async def index_order_in_milvus(order: Dict, log: bool = True) -> int:
    """Index a newly created order into Milvus as retrieval data for the LLM.

    Returns the number of vectors indexed (header + items).
    """
    try:
        # Milvus optional; ignore ensure step when disabled or inaccessible
        ensure_collection_exists()
    except Exception:
        pass
    try:
        order_id = order.get("id") or ""
        source = order.get("source_file", "order")
        items = order.get("items", []) or []
        subtotal = float(order.get("subtotal", 0) or 0)
        status = order.get("status") or "Quoted"
        header = (
            f"Order {order_id} | source: {source} | status: {status} | "
            f"items: {len(items)} | subtotal: ${subtotal:,.0f}"
        )
        tasks = [add_document(header, order_id, 0, source)]
        for idx, it in enumerate(items, start=1):
            name = str(it.get("item", "")).strip()
            try:
                qty = float(it.get("quantity", 0) or 0)
            except Exception:
                qty = 0.0
            try:
                unit = float(it.get("unit_cost", it.get("cost", 0)) or 0)
            except Exception:
                unit = 0.0
            line_total = float(it.get("line_total", qty * unit) or 0)
            line = f"{name} | qty {qty:g} | unit ${unit:,.0f} | total ${line_total:,.0f}"
            tasks.append(add_document(line, order_id, idx, source))
        await asyncio.gather(*tasks, return_exceptions=True)
        if log:
            try:
                append_log({
                    "type": "order_indexed_milvus",
                    "order_id": order_id,
                    "details": {"items_indexed": len(tasks)},
                })
            except Exception:
                pass
        return int(len(tasks))
    except Exception:
        return 0

async def index_notes_in_milvus(order_id: str, notes: str, source: str = "notetaker", log: bool = True) -> int:
    """Chunk and index freeform notes/transcripts so the agent can retrieve evidence later.

    Returns the number of note chunks indexed.
    """
    if not notes or not order_id:
        return 0
    try:
        # Milvus optional; ignore ensure step when disabled or inaccessible
        ensure_collection_exists()
    except Exception:
        pass
    try:
        chunks = _chunk_text(notes, chunk_size=800, chunk_overlap=120)
        tasks = [add_document(chunk, f"{order_id}:notes", idx, source) for idx, chunk in enumerate(chunks)]
        await asyncio.gather(*tasks, return_exceptions=True)
        if log:
            try:
                append_log({
                    "type": "notes_indexed_milvus",
                    "order_id": order_id,
                    "details": {"chunks": len(chunks)},
                })
            except Exception:
                pass
        return int(len(chunks))
    except Exception:
        return 0


async def index_all_in_milvus(order: Dict, notes: Optional[str] = None) -> None:
    """Index order header/items and optional notes, then emit a single combined log."""
    try:
        order_id = order.get("id") or ""
        # Run both indexing tasks concurrently without their own logs
        tasks: list = [index_order_in_milvus(order, log=False)]
        if notes:
            tasks.append(index_notes_in_milvus(order_id, notes, log=False))
        results = await asyncio.gather(*tasks, return_exceptions=True)
        items_indexed = 0
        note_chunks = 0
        try:
            items_indexed = int(results[0]) if isinstance(results[0], (int, float)) else 0
        except Exception:
            items_indexed = 0
        if len(results) > 1:
            try:
                note_chunks = int(results[1]) if isinstance(results[1], (int, float)) else 0
            except Exception:
                note_chunks = 0
        append_log({
            "type": "ingest_indexed_milvus",
            "order_id": order_id,
            "details": {"items_indexed": items_indexed, "note_chunks": note_chunks},
        })
    except Exception:
        # Best-effort logging only
        pass

@app.post("/agent/upload", operation_id="agent_upload") #this is the pdf quote that is auto-generated you can download from orders page
async def agent_upload(file: UploadFile = File(...)):
    """Upload .pdf or .docx, chunk, and insert into Milvus."""
    content = await file.read()
    text = ""
    try:
        if file.filename.lower().endswith(".docx"):
            if not Document:
                return {"error": "python-docx not installed"}
            doc = Document(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        elif file.filename.lower().endswith(".pdf"):
            if not PdfReader:
                return {"error": "pypdf not installed"}
            reader = PdfReader(BytesIO(content))
            pages = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t)
            text = "\n".join(pages)
        else:
            return {"error": "Unsupported file type. Upload .pdf or .docx"}
    except Exception as e:
        return {"error": f"Failed to read file: {e}"}

    if not text.strip():
        return {"error": "No extractable text found in the file"}

    try:
        ensure_collection_exists()
    except Exception:
        pass

    chunks = _chunk_text(text)
    doc_id = Path(file.filename).stem
    source = file.filename
    tasks = [add_document(chunk, doc_id, idx, source) for idx, chunk in enumerate(chunks)]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    success = sum(1 for r in results if isinstance(r, dict) and r.get("message", "").lower().startswith("document"))

    # Extract quote-like line items and compute subtotal
    items = await extract_line_items(text)
    subtotal = 0.0
    normalized_items = []
    for it in items:
        name = str(it.get("item", "")).strip()
        try:
            qty = float(it.get("quantity", 0))
        except Exception:
            qty = 0.0
        try:
            unit_cost = float(it.get("unit_cost", it.get("cost", 0)))
        except Exception:
            unit_cost = 0.0
        line_total = qty * unit_cost
        subtotal += line_total
        normalized_items.append({
            "item": name,
            "quantity": qty,
            "unit_cost": unit_cost,
            "line_total": line_total,
        })

    # Try to extract a customer name from the document text
    try:
        extracted_customer = _extract_customer_from_notes(text)
        if extracted_customer:
            extracted_customer = _canonicalize_customer(extracted_customer)
    except Exception:
        extracted_customer = None
    # Include prompt context if upload originated from an agent-run flow in the same session
    order = create_order(
        file.filename,
        subtotal,
        len(normalized_items),
        status="Quoted",
        items=normalized_items,
        prompt=_summarize_text((AGENT_PROMPT.get() or recent_user_prompt()), 200),
        tool_name="agent_upload",
    )
    # If we found a customer, update immediately so Pricing sees it on first load
    try:
        if extracted_customer:
            _ = update_customer(order.get("id"), extracted_customer, prompt=(AGENT_PROMPT.get() or recent_user_prompt()), tool_name="agent_upload")
    except Exception:
        pass
    append_log({
        "type": "upload_ingest",
        "file": file.filename,
        "chunks": len(chunks),
        "inserted": success,
        "order_id": order.get("id"),
        "items_count": len(normalized_items),
        "subtotal": subtotal,
        "prompt": (_summarize_text(recent_user_prompt(), 200) or _summarize_text(AGENT_PROMPT.get(), 200)),
        "tool_name": "agent_upload",
        "details": {"mcp_client_triggered": True, "customer": (extracted_customer or "")},
    })
    try:
        asyncio.create_task(index_all_in_milvus(order, notes=None))
    except Exception:
        pass
    # Kick off MCP client asynchronously for post-upload quote generation (combined log above)
    try:
        asyncio.create_task(mcp_client_run())
    except Exception as e:
        append_log({"type": "mcp_client_error", "order_id": order.get("id"), "error": str(e)})
    return {
        "file": file.filename,
        "chunks": len(chunks),
        "inserted": success,
        "items": normalized_items,
        "subtotal": subtotal,
        "order": order,
        "pricing_url": f"/pricing/{order.get('id')}"
    }


async def extract_line_items(text: str):
    """Extract structured line items from text using the OpenAI client.

    Returns: List[{item, quantity, unit_cost}] (best-effort)
    """
    from src.milvus_connector import client as openai_client
    system_msg = (
        "Extract a list of quote line items from the text. "
        "Respond ONLY with JSON in the form: {\"items\":[{\"item\":str,\"quantity\":number,\"unit_cost\":number}, ...]} . "
        "If values are missing, infer reasonable quantities and set unit_cost to 0."
    )
    try:
        resp = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": text[:20000]},
            ],
            temperature=0.2,
            max_tokens=800,
            stream=False,
        )
        content = resp.choices[0].message.content.strip()
        import json as pyjson, re
        try:
            data = pyjson.loads(content)
        except Exception:
            m = re.search(r"\{[\s\S]*\}", content)
            data = pyjson.loads(m.group(0)) if m else {"items": []}
        items = data.get("items", [])
        return items if isinstance(items, list) else []
    except Exception:
        return []


@app.post("/agent/quote_pdf", operation_id="agent_quote_pdf")
async def agent_quote_pdf(payload: Dict, request: Request):
    """Generate a PDF from provided items and subtotal; return inline for a new tab."""
    if canvas is None:
        return JSONResponse(status_code=400, content={"error": "reportlab not installed"})
    items = payload.get("items", [])
    try:
        subtotal = float(payload.get("subtotal", 0))
    except Exception:
        subtotal = 0.0

    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    y = height - 1 * inch

    c.setFont("Helvetica-Bold", 16)
    c.drawString(1 * inch, y, "Quote Preview")
    y -= 0.4 * inch
    c.setFont("Helvetica", 10)
    c.drawString(1 * inch, y, "Item")
    c.drawString(4.2 * inch, y, "Qty")
    c.drawString(4.8 * inch, y, "Unit Cost")
    c.drawString(5.8 * inch, y, "Line Total")
    y -= 0.25 * inch

    c.setFont("Helvetica", 9)
    for it in items:
        if y < 1 * inch:
            c.showPage()
            y = height - 1 * inch
            c.setFont("Helvetica", 9)
        name = str(it.get("item", ""))[:60]
        try:
            qty = float(it.get("quantity", 0))
        except Exception:
            qty = 0.0
        try:
            unit_cost = float(it.get("unit_cost", 0))
        except Exception:
            unit_cost = 0.0
        line_total = float(it.get("line_total", qty * unit_cost))
        c.drawString(1 * inch, y, name)
        c.drawRightString(4.6 * inch, y, f"{qty:g}")
        c.drawRightString(5.6 * inch, y, f"${unit_cost:,.2f}")
        c.drawRightString(7.5 * inch, y, f"${line_total:,.2f}")
        y -= 0.22 * inch

    y -= 0.2 * inch
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(7.5 * inch, y, f"Subtotal: ${subtotal:,.2f}")

    c.showPage()
    c.save()
    pdf_buffer.seek(0)
    headers = {
        "Content-Disposition": "inline; filename=quote_preview.pdf",
    }
    # If an MCP client (or any JSON-preferring client) calls this tool, return a JSON stub
    # instead of a binary PDF to avoid decode errors in tool adapters.
    try:
        accept = str(request.headers.get("accept", "")).lower()
        if "application/json" in accept:
            append_log({
                "type": "quote_pdf",
                "items": len(items),
                "subtotal": subtotal,
                "prompt": _audit_prompt_context(),
            })
            return JSONResponse({
                "message": "PDF generated for preview",
                "note": "Binary PDF not returned in JSON mode. Call this endpoint from a browser to stream the PDF inline.",
            })
    except Exception:
        pass
    append_log({
        "type": "quote_pdf",
        "items": len(items),
        "subtotal": subtotal,
        "prompt": (_summarize_text(recent_user_prompt(), 200) or _summarize_text(AGENT_PROMPT.get(), 200)),
    })
    return StreamingResponse(pdf_buffer, media_type="application/pdf", headers=headers)


@app.get("/orders/pdf/{order_id}", operation_id="orders_pdf")
async def orders_pdf(order_id: str, request: Request):
    if canvas is None:
        return JSONResponse(status_code=400, content={"error": "reportlab not installed"})
    order = get_order(order_id)
    if not order:
        return JSONResponse(status_code=404, content={"error": "order not found"})
    items = order.get("items", [])
    subtotal = float(order.get("subtotal", 0))

    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    y = height - 1 * inch

    c.setFont("Helvetica-Bold", 16)
    c.drawString(1 * inch, y, f"Quote {order_id}")
    y -= 0.4 * inch
    c.setFont("Helvetica", 10)
    c.drawString(1 * inch, y, f"Customer: {order.get('customer','')}")
    y -= 0.2 * inch
    c.drawString(1 * inch, y, f"Source: {order.get('source_file','')}")
    y -= 0.25 * inch
    c.drawString(1 * inch, y, "Item")
    c.drawString(4.2 * inch, y, "Qty")
    c.drawString(4.8 * inch, y, "Unit Cost")
    c.drawString(5.8 * inch, y, "Line Total")
    y -= 0.25 * inch

    c.setFont("Helvetica", 9)
    for it in items:
        if y < 1 * inch:
            c.showPage()
            y = height - 1 * inch
            c.setFont("Helvetica", 9)
        name = str(it.get("item", ""))[:60]
        try:
            qty = float(it.get("quantity", 0))
        except Exception:
            qty = 0.0
        try:
            unit_cost = float(it.get("unit_cost", 0))
        except Exception:
            unit_cost = 0.0
        line_total = float(it.get("line_total", qty * unit_cost))
        c.drawString(1 * inch, y, name)
        c.drawRightString(4.6 * inch, y, f"{qty:g}")
        c.drawRightString(5.6 * inch, y, f"${unit_cost:,.2f}")
        c.drawRightString(7.5 * inch, y, f"${line_total:,.2f}")
        y -= 0.22 * inch

    y -= 0.2 * inch
    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(7.5 * inch, y, f"Subtotal: ${subtotal:,.2f}")

    c.showPage()
    c.save()
    pdf_buffer.seek(0)
    # If called by MCP (expects JSON), return a JSON descriptor rather than raw PDF
    try:
        accept = str(request.headers.get("accept", "")).lower()
        if "application/json" in accept:
            return JSONResponse({
                "pdf_url": f"/orders/pdf/{order_id}",
                "message": "Open the URL in a browser to view the PDF",
            })
    except Exception:
        pass
    headers = {"Content-Disposition": f"inline; filename={order_id}.pdf"}
    return StreamingResponse(pdf_buffer, media_type="application/pdf", headers=headers)


@app.get("/logs/data", operation_id="logs_data")
async def logs_data(limit: int = 20, offset: int = 0):
    return {"logs": list_logs(limit=limit, offset=offset)}


class LogEmit(BaseModel):
    type: str
    order_id: str | None = None
    details: dict | None = None
    prompt: str | None = None


@app.post("/logs/emit", operation_id="logs_emit")
async def logs_emit(body: LogEmit):
    try:
        rec = append_log({
            "type": body.type,
            "order_id": body.order_id,
            "details": body.details,
            # Summarize prompt if provided; else use recent user prompt
            "prompt": (_summarize_text(body.prompt, 200) if body.prompt else _summarize_text(recent_user_prompt(), 200)),
        })
        # If this is an assistant message, remember it for short-confirmation context
        try:
            if str(body.type) == "assistant_message" and isinstance(body.details, dict):
                msg = body.details.get("message")
                remember_assistant_message(msg)
        except Exception:
            pass
        return {"ok": True, "log": rec}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


class NotetakerIngest(BaseModel):
    notes: str = Field(..., description="Freeform transcript or call notes to generate a quote from")


@app.post("/notetaker/ingest", operation_id="notetaker_ingest")
async def notetaker_ingest(body: NotetakerIngest, _auth=Depends(require_auth)):
    """Integration endpoint: ingest call notes, let the agent act, and return the new order id and pricing URL."""
    notes = body.notes or ""
    try:
        append_log({
            "type": "notetaker_ingest_start",
            "prompt": _summarize_text(notes, 200),
            "tool_name": "notetaker_ingest",
            "user_id": (_auth or {}).get("sub") if isinstance(_auth, dict) else None,
        })
        # Parse items from notes and create an order deterministically (no agent round-trip)
        items = await extract_line_items(notes)
        subtotal = 0.0
        normalized_items = []
        for it in items or []:
            name = str(it.get("item", "")).strip()
            try:
                qty = float(it.get("quantity", 0))
            except Exception:
                qty = 0.0
            try:
                unit_cost = float(it.get("unit_cost", it.get("cost", 0)))
            except Exception:
                unit_cost = 0.0
            line_total = qty * unit_cost
            subtotal += line_total
            normalized_items.append({
                "item": name,
                "quantity": qty,
                "unit_cost": unit_cost,
                "line_total": line_total,
            })

        order = create_order(
            source_file="Call Notes",
            subtotal=subtotal,
            items_count=len(normalized_items),
            status="Quoted",
            items=normalized_items,
            prompt=notes,
            tool_name="notetaker_ingest",
        )

        order_id = order.get("id")
        # Extract and set customer name if present in notes
        try:
            extracted = _extract_customer_from_notes(notes)
            if extracted:
                extracted = _canonicalize_customer(extracted)
            if extracted:
                _ = update_customer(order_id, extracted, prompt=notes, tool_name="notetaker_ingest")
        except Exception:
            pass
        # Index order and notes in Milvus
        try:
            asyncio.create_task(index_all_in_milvus(order, notes))
        except Exception:
            pass

        append_log({
            "type": "notetaker_ingest",
            "order_id": order_id,
            "prompt": _summarize_text(notes, 200),
            "tool_name": "notetaker_ingest",
            "user_id": (_auth or {}).get("sub") if isinstance(_auth, dict) else None,
            "details": {
                "items_count": len(normalized_items),
                "subtotal": subtotal,
                "customer": extracted or "",
                "mcp_client_triggered": True,
            },
        })
        return {"ok": True, "order_id": order_id, "pricing_url": f"/pricing/{order_id}"}
    except Exception as e:
        append_log({
            "type": "notetaker_ingest_error",
            "prompt": _summarize_text(notes, 200),
            "tool_name": "notetaker_ingest",
            "details": {"error": str(e)},
            "user_id": (_auth or {}).get("sub") if isinstance(_auth, dict) else None,
        })
        return JSONResponse(status_code=400, content={"error": str(e)})


class PricingUpdate(BaseModel):
    items: list
    prompt: Optional[str] = None


@app.get("/pricing/{order_id}", operation_id="pricing_page")
async def pricing_page(order_id: str):
    page = static_dir / "pricing.html"
    if page.exists():
        return FileResponse(str(page), media_type="text/html")
    return JSONResponse(status_code=404, content={"error": "pricing.html not found"})


@app.get("/pricing/data/{order_id}", operation_id="pricing_data")
async def pricing_data(order_id: str):
    order = get_order(order_id)
    if not order:
        return JSONResponse(status_code=404, content={"error": "order not found"})
    # Fallback: if items are missing but items_count suggests there should be items, try local cache
    items = order.get("items", []) or []
    try:
        has_items_count = float(order.get("items_count", 0) or 0) > 0
    except Exception:
        has_items_count = False
    if (not items) and has_items_count:
        try:
            data_path = Path(__file__).resolve().parents[1] / "data" / "orders.json"
            if data_path.exists():
                import json as _json
                store = _json.loads(data_path.read_text(encoding="utf-8"))
                for o in (store.get("orders", []) or []):
                    if o.get("id") == order_id and o.get("items"):
                        items = o.get("items")
                        break
            # Best-effort: persist back if items found
            if items:
                try:
                    _ = update_items(order_id, items, prompt=AGENT_PROMPT.get(), tool_name="pricing_data")
                except Exception:
                    pass
        except Exception:
            pass
    # Recompute subtotal from items if needed
    try:
        if items and not order.get("subtotal"):
            subtotal = 0.0
            for it in items:
                q = float(it.get("quantity", 0) or 0)
                c = float(it.get("unit_cost", it.get("cost", 0)) or 0)
                subtotal += q * c
            order["subtotal"] = subtotal
    except Exception:
        pass
    return {
        "order": {
            "id": order.get("id"),
            "source_file": order.get("source_file"),
            "customer": order.get("customer", ""),
            "items": items,
            "subtotal": order.get("subtotal", 0),
            "status": order.get("status")
        }
    }


@app.post("/pricing/save/{order_id}", operation_id="pricing_save")
async def pricing_save(order_id: str, body: PricingUpdate, _auth=Depends(require_auth)):
    effective_prompt = body.prompt if getattr(body, "prompt", None) else (AGENT_PROMPT.get() or recent_agent_prompt())
    updated = update_items(order_id, (body.items or []), prompt=effective_prompt, tool_name="pricing_save")
    if not updated:
        return JSONResponse(status_code=404, content={"error": "order not found"})
    return {"ok": True, "order": updated}


@app.get(
    "/orders/data",
    operation_id="list_orders",
    summary="List all orders",
    description="Returns an object with `orders`, where each order contains: `id`, `subtotal`, `status`, `source_file`, `items_count`, `created_at` (and optionally `items`)."
)
async def orders_data():
    return {"orders": list_orders()}


class OrderStatusUpdate(BaseModel):
    id: str
    status: str
    prompt: Optional[str] = None


class OrdersStatusBulkUpdate(BaseModel):
    ids: list[str]
    status: str
    prompt: Optional[str] = None


@app.post(
    "/orders/status",
    operation_id="update_order_status",
    summary="Update a single order status",
    description="Body must be `{ id: string, status: Quoted|Sent|Received }`. Returns the updated order."
)
async def orders_status(update: OrderStatusUpdate, _auth=Depends(require_auth)):
    if update.status not in ("Quoted", "Sent", "Received"):
        return JSONResponse(status_code=400, content={"error": "invalid status"})
    # Use explicit prompt if provided; else fallback to current/last agent prompt context
    effective_prompt = update.prompt if update.prompt else (AGENT_PROMPT.get() or recent_agent_prompt())
    effective_id = _resolve_order_id(update.id) or update.id
    updated = update_status(effective_id, update.status, prompt=effective_prompt, tool_name="update_order_status")
    if not updated:
        return JSONResponse(status_code=404, content={"error": "order not found"})
    try:
        append_log({
            "type": "order_status_updated",
            "order_id": effective_id,
            "details": {"status": update.status},
            "prompt": _audit_prompt_context(effective_prompt),
            "tool_name": "update_order_status",
        })
    except Exception:
        pass
    return updated


@app.post(
    "/orders/status/bulk",
    operation_id="update_orders_status_bulk",
    summary="Update status for multiple orders",
    description="Body must be `{ ids: string[], status: Quoted|Sent|Received }`. Returns `{ updated_ids: string[] }`."
)
async def orders_status_bulk(update: OrdersStatusBulkUpdate, _auth=Depends(require_auth)):
    if update.status not in ("Quoted", "Sent", "Received"):
        return JSONResponse(status_code=400, content={"error": "invalid status"})
    ids = list(update.ids or [])
    if not ids:
        return JSONResponse(status_code=400, content={"error": "missing ids"})
    effective_prompt = update.prompt if update.prompt else (AGENT_PROMPT.get() or recent_agent_prompt())
    updated_ids: list[str] = []
    for raw_id in ids:
        effective_id = _resolve_order_id(raw_id) or raw_id
        try:
            res = update_status(effective_id, update.status, prompt=effective_prompt, tool_name="update_orders_status_bulk")
            if res:
                updated_ids.append(effective_id)
        except Exception:
            # continue best-effort
            pass
    try:
        append_log({
            "type": "orders_status_updated_bulk",
            "details": {"status": update.status, "updated_ids": updated_ids},
            "prompt": _audit_prompt_context(effective_prompt),
            "tool_name": "update_orders_status_bulk",
        })
    except Exception:
        pass
    return {"updated_ids": updated_ids}
class OrderCustomerUpdate(BaseModel):
    id: Optional[str] = None
    customer: str
    prompt: Optional[str] = None


@app.post(
    "/orders/customer",
    operation_id="update_order_customer",
    summary="Update order customer name",
    description="Body must be `{ id: string, customer: string }`. Returns the updated order."
)
async def orders_customer(update: OrderCustomerUpdate, _auth=Depends(require_auth)):
    # Fallback to current order context if id not provided
    raw_id = update.id or (AGENT_ORDER_ID.get() or recent_agent_order_id())
    effective_id = _resolve_order_id(raw_id) or raw_id
    if not effective_id:
        return JSONResponse(status_code=400, content={"error": "missing order id"})
    effective_prompt = update.prompt if update.prompt else (AGENT_PROMPT.get() or recent_agent_prompt())
    # Ensure order is mutable; if Sent/Received, set back to Quoted before customer change
    try:
        current = get_order(effective_id)
        if current and str(current.get("status")) in ("Sent", "Received"):
            _ = update_status(effective_id, "Quoted", prompt=effective_prompt, tool_name="update_order_customer")
    except Exception:
        pass
    updated = update_customer(effective_id, update.customer, prompt=effective_prompt, tool_name="update_order_customer")
    if not updated:
        # As a last resort, check if order exists to return a clearer error
        exists = get_order(effective_id)
        return JSONResponse(status_code=404, content={"error": "order not found" if not exists else "failed to update customer"})
    try:
        append_log({
            "type": "order_customer_updated",
            "order_id": effective_id,
            "details": {"customer": update.customer},
            "prompt": _audit_prompt_context(effective_prompt),
            "tool_name": "update_order_customer",
        })
    except Exception:
        pass
    return updated


@app.get("/orders/stats")
async def orders_stats():
    orders = list_orders()
    total_cost = sum(float(o.get("subtotal", 0) or 0) for o in orders)
    open_count = sum(1 for o in orders if (o.get("status") in ("Quoted", "Sent")))
    return {
        "total_cost": total_cost,
        "open_count": open_count,
        "total_count": len(orders)
    }


@app.get("/orders/top", operation_id="identify_orders")
async def orders_top(limit: int = 3, include_status: str = "any") -> Dict:
    """Identify highest-cost orders (for MCP usage). Optionally filter by status.

    - limit: number of orders to return (default 3)
    - include_status: one of "any", "open" (Quoted/Sent), or a specific status
    """
    all_orders = list_orders()
    if include_status == "open":
        candidates = [o for o in all_orders if o.get("status") in ("Quoted", "Sent")]
    elif include_status in ("Quoted", "Sent", "Received"):
        candidates = [o for o in all_orders if o.get("status") == include_status]
    else:
        candidates = all_orders

    # Sort by subtotal descending and take top N
    top = sorted(candidates, key=lambda o: float(o.get("subtotal", 0) or 0), reverse=True)[: max(1, limit)]

    # Minimal projection for LLM consumption
    result = [
        {
            "id": o.get("id"),
            "subtotal": float(o.get("subtotal", 0) or 0),
            "status": o.get("status"),
            "source_file": o.get("source_file"),
        }
        for o in top
    ]

    # Audit trail (GET)
    append_log({
        "type": "orders_identified_top",
        "details": {"limit": limit, "include_status": include_status, "ids": [r["id"] for r in result]},
        "route": "/orders/top",
        "method": "GET",
        "status": 200,
        "prompt": _audit_prompt_context(),
        "tool_name": "identify_orders",
    })
    return {"orders": result}


class CloneOrderRequest(BaseModel):
    source_order_id: str
    additions: list = Field(default_factory=list, description="List of {item, quantity, unit_cost}")
    prompt: Optional[str] = None


def _to_float(value) -> float:
    try:
        if isinstance(value, (int, float)):
            return float(value)
        s = str(value).strip().replace("$", "").replace(",", "")
        return float(s) if s else 0.0
    except Exception:
        return 0.0


def _normalize_items(items: list) -> tuple[list, float]:
    normalized = []
    subtotal = 0.0
    for it in items or []:
        name = str(it.get("item", "")).strip()
        qty_raw = it.get("quantity", it.get("qty", it.get("q", it.get("count", 0))))
        cost_raw = it.get(
            "unit_cost",
            it.get(
                "cost",
                it.get("amount", it.get("price", it.get("unitPrice", it.get("unit_price", 0)))),
            ),
        )
        qty = _to_float(qty_raw)
        unit_cost = _to_float(cost_raw)
        line_total = qty * unit_cost
        subtotal += line_total
        normalized.append(
            {
                "item": name,
                "quantity": qty,
                "unit_cost": unit_cost,
                "line_total": line_total,
            }
        )
    return normalized, float(subtotal)


@app.post("/orders/clone", operation_id="clone_order")
async def clone_order(payload: CloneOrderRequest, _auth=Depends(require_auth)) -> Dict:
    """Clone an existing order and add line items (e.g., Service Charge). Returns the new order."""
    try:
        source = get_order(payload.source_order_id)
        if not source:
            return JSONResponse(status_code=404, content={"error": "source order not found"})

        base_items = source.get("items", [])
        add_items, _ = _normalize_items(payload.additions)
        all_items = base_items + add_items
        normalized, subtotal = _normalize_items(all_items)

        source_name = source.get("source_file") or f"Cloned from {payload.source_order_id}"
        effective_prompt = payload.prompt if getattr(payload, "prompt", None) else (AGENT_PROMPT.get() or recent_agent_prompt())
        new_order = create_order(source_name, subtotal, len(normalized), status="Quoted", items=normalized, prompt=effective_prompt, tool_name="clone_order")

        append_log(
            {
                "type": "order_cloned",
                "order_id": new_order.get("id"),
                "details": {
                    "source": payload.source_order_id,
                    "source_file": source_name,
                    "additions": payload.additions,
                    "items_count": len(normalized),
                    "subtotal": subtotal,
                },
                "route": "/orders/clone",
                "method": "POST",
                "status": 200,
                # Use last user prompt summary instead of tool guide
                "prompt": _audit_prompt_context(effective_prompt),
                "tool_name": "clone_order",
            }
        )
        try:
            asyncio.create_task(index_order_in_milvus(new_order))
        except Exception:
            pass
        return {"order": new_order}
    except Exception as e:
        append_log({"type": "order_clone_error", "details": {"error": str(e), "payload": payload.dict()}})
        return JSONResponse(status_code=400, content={"error": str(e)})


class AddItemsRequest(BaseModel):
    order_id: str
    additions: list = Field(default_factory=list, description="List of {item, quantity, unit_cost}")
    prompt: Optional[str] = None


@app.post("/orders/add_items", operation_id="add_items_to_order")
async def add_items_to_order(payload: AddItemsRequest, _auth=Depends(require_auth)) -> Dict:
    """Add line items to an existing order (two-step flows: clone → add items)."""
    try:
        existing = get_order(payload.order_id)
        if not existing:
            return JSONResponse(status_code=404, content={"error": "order not found"})
        base_items = list(existing.get("items") or [])
        add_items, _ = _normalize_items(payload.additions)
        merged = base_items + add_items
        effective_prompt = payload.prompt if getattr(payload, "prompt", None) else (AGENT_PROMPT.get() or recent_agent_prompt())
        updated = update_items(payload.order_id, merged, prompt=effective_prompt, tool_name="add_items_to_order")
        if not updated:
            return JSONResponse(status_code=400, content={"error": "failed to update items"})
        append_log({
            "type": "order_items_added",
            "order_id": payload.order_id,
            "details": {
                "added": add_items,
                "new_items_count": len(updated.get("items", [])),
                "items_count": len(updated.get("items", [])),
                "subtotal": updated.get("subtotal", 0)
            },
            "prompt": _audit_prompt_context(effective_prompt),
            "tool_name": "add_items_to_order",
        })
        return {"order": updated}
    except Exception as e:
        append_log({"type": "order_items_add_error", "details": {"error": str(e), "payload": payload.dict()}})
        return JSONResponse(status_code=400, content={"error": str(e)})

class CreateVariantRequest(BaseModel):
    source_order_id: str = Field(..., description="Existing order id to clone from")
    additions: list = Field(default_factory=list, description="List of {item, quantity, unit_cost}")
    status: str = Field(default="Quoted", description="Status for the new order (Quoted|Sent|Received)")
    prompt: Optional[str] = None


@app.post("/orders/variant", operation_id="create_variant_order")
async def create_variant_order(payload: CreateVariantRequest, _auth=Depends(require_auth)) -> Dict:
    """Create a new order variant from an existing one, with added items and desired status.

    This endpoint NEVER modifies the source order. It returns the newly created order id.
    """
    try:
        if payload.status not in ("Quoted", "Sent", "Received"):
            return JSONResponse(status_code=400, content={"error": "invalid status"})
        # Be forgiving about whitespace in incoming ids
        source_id = (payload.source_order_id or "").strip()
        source = get_order(source_id)
        if not source:
            return JSONResponse(status_code=404, content={"error": "source order not found"})

        base_items = source.get("items", [])
        add_items, _ = _normalize_items(payload.additions)
        combined = base_items + add_items
        normalized, subtotal = _normalize_items(combined)

        source_name = source.get("source_file") or f"Variant of {source_id}"
        effective_prompt = payload.prompt if getattr(payload, "prompt", None) else (AGENT_PROMPT.get() or recent_agent_prompt())
        new_order = create_order(source_name, subtotal, len(normalized), status=payload.status, items=normalized, prompt=effective_prompt, tool_name="create_variant_order")

        append_log(
            {
                "type": "order_variant_created",
                "order_id": new_order.get("id"),
                "details": {
                    "source": payload.source_order_id,
                    "additions": payload.additions,
                    "status": payload.status,
                    "items_count": len(normalized),
                    "subtotal": subtotal,
                },
                "route": "/orders/variant",
                "method": "POST",
                "status": 200,
                "prompt": _audit_prompt_context(effective_prompt),
                "tool_name": "create_variant_order",
            }
        )
        try:
            asyncio.create_task(index_order_in_milvus(new_order))
        except Exception:
            pass
        return {"order": new_order}
    except Exception as e:
        append_log({"type": "order_variant_error", "details": {"error": str(e), "payload": payload.dict()}})
        return JSONResponse(status_code=400, content={"error": str(e)})

# Initialize FastAPIMCP AFTER all routes are defined so every endpoint is exposed as a tool
mcp_server = FastApiMCP(
    app,
    name="MCP API",
    description="MCP server for the API",
    describe_full_response_schema=True,
    describe_all_responses=True,
)
mcp_server.mount()

# Compatibility shim: some clients probe POST /mcp before establishing message channel.
# Return Method Not Allowed as JSON instead of framework default HTML for faster fail & clearer logs.
@app.post("/mcp")
async def _mcp_probe_shim():
    return JSONResponse(status_code=405, content={"error": "Method Not Allowed. Use GET /mcp then POST /mcp/messages."})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 